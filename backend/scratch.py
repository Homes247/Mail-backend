import sys
import io
import docx
import base64

def test_parse(filepath):
    try:
        doc_file = docx.Document(filepath)
        html = ""

        def _run_to_html(run):
            text = run.text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            if not text:
                return ""
            if run.bold:      text = f"<b>{text}</b>"
            if run.italic:    text = f"<i>{text}</i>"
            if run.underline: text = f"<u>{text}</u>"
            if run.font.strike: text = f"<s>{text}</s>"
            return text

        def _alignment_style(p):
            try:
                if p.alignment == 1: return ' style="text-align:center;"'
                if p.alignment == 2: return ' style="text-align:right;"'
                if p.alignment == 3: return ' style="text-align:justify;"'
            except:
                pass
            return ""

        def _para_tag(p):
            try:
                style = (p.style.name or "").lower()
                if "heading 1" in style: return "h1"
                if "heading 2" in style: return "h2"
                if "heading 3" in style: return "h3"
                if "heading 4" in style: return "h4"
            except:
                pass
            return "p"

        for element in doc_file.element.body:
            tag = element.tag.split("}")[-1]  # strip namespace

            if tag == "p":
                from docx.text.paragraph import Paragraph
                # IS THIS CORRECT?
                p = Paragraph(element, doc_file._body)
                p_html = ""

                for run in p.runs:
                    p_html += _run_to_html(run)

                if p_html.strip():
                    tag_name = _para_tag(p)
                    html += f"<{tag_name}{_alignment_style(p)}>{p_html}</{tag_name}>"
                elif not p_html:
                    html += "<br>"

            elif tag == "tbl":
                from docx.table import Table
                tbl = Table(element, doc_file._body)
                html += '<table border="1" style="border-collapse:collapse;width:100%;margin:8px 0;"><tbody>'
                for row in tbl.rows:
                    html += "<tr>"
                    for cell in row.cells:
                        cell_text = cell.text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                        html += f'<td style="padding:4px;">{cell_text}</td>'
                    html += "</tr>"
                html += "</tbody></table>"

        print("HTML LENGTH:", len(html))
        print(html[:200])
    except Exception as e:
        import traceback
        traceback.print_exc()

# create a dummy docx
doc = docx.Document()
doc.add_paragraph("Hello world")
doc.save("test.docx")
test_parse("test.docx")
