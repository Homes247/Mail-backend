import sys
import io
import docx
import base64

def test_parse(filepath):
    try:
        doc_file = docx.Document(filepath)
        for element in doc_file.element.body:
            tag = element.tag.split("}")[-1]  # strip namespace
            if tag == "p":
                from docx.text.paragraph import Paragraph
                # This is what I used in documents.py
                p = Paragraph(element, doc_file)
                for run in p.runs:
                    print(run.text)
    except Exception as e:
        import traceback
        traceback.print_exc()

doc = docx.Document()
doc.add_paragraph("Hello world")
doc.save("test2.docx")
test_parse("test2.docx")
