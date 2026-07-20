import sys
import io
import docx
import base64

def test_parse(filepath):
    try:
        doc_file = docx.Document(filepath)
        for element in doc_file.element.body:
            tag = element.tag.split("}")[-1]  # strip namespace
            if tag == "tbl":
                from docx.table import Table
                # This is what I used in documents.py
                t = Table(element, doc_file)
                for row in t.rows:
                    for cell in row.cells:
                        print(cell.text)
    except Exception as e:
        import traceback
        traceback.print_exc()

doc = docx.Document()
doc.add_table(rows=1, cols=1)
doc.tables[0].cell(0, 0).text = "Test Table"
doc.save("test3.docx")
test_parse("test3.docx")
